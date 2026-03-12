from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict, Any
from loguru import logger
import os

class CAMGenerator:
    """Generate professional Credit Appraisal Memo documents"""
    
    def __init__(self, output_dir: str = "./outputs"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def generate_cam(self, 
                     company_info: Dict[str, Any],
                     data_summary: Dict[str, Any],
                     research_summary: Dict[str, Any],
                     recommendation: Dict[str, Any],
                     primary_insights: Dict[str, Any] = None) -> str:
        """Generate comprehensive CAM document"""
        
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Title
        title = doc.add_heading('CREDIT APPRAISAL MEMO', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Company header
        doc.add_heading(f"{company_info.get('company_name', 'N/A')}", 1)
        
        # Executive Summary
        doc.add_heading('EXECUTIVE SUMMARY', 2)
        self._add_executive_summary(doc, company_info, recommendation)
        
        # Recommendation
        doc.add_heading('CREDIT RECOMMENDATION', 2)
        self._add_recommendation_section(doc, recommendation)
        
        # Five Cs Analysis
        doc.add_heading('FIVE Cs OF CREDIT ANALYSIS', 2)
        self._add_five_cs_analysis(doc, recommendation)
        
        # Company Overview
        doc.add_heading('COMPANY OVERVIEW', 2)
        self._add_company_overview(doc, company_info, data_summary)
        
        # Financial Analysis
        doc.add_heading('FINANCIAL ANALYSIS', 2)
        self._add_financial_analysis(doc, data_summary)
        
        # GST & Banking Analysis
        doc.add_heading('GST & BANKING ANALYSIS', 2)
        self._add_gst_banking_analysis(doc, data_summary)
        
        # Secondary Research Findings
        doc.add_heading('SECONDARY RESEARCH FINDINGS', 2)
        self._add_research_findings(doc, research_summary)
        
        # Primary Due Diligence
        if primary_insights:
            doc.add_heading('PRIMARY DUE DILIGENCE', 2)
            self._add_primary_insights(doc, primary_insights)
        
        # Risk Assessment
        doc.add_heading('RISK ASSESSMENT', 2)
        self._add_risk_assessment(doc, recommendation)
        
        # Terms & Conditions
        doc.add_heading('PROPOSED TERMS & CONDITIONS', 2)
        self._add_terms_conditions(doc, recommendation)
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d-%b-%Y %H:%M')}")
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Save document
        filename = f"CAM_{company_info.get('company_name', 'Company').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        logger.info(f"CAM generated: {filepath}")
        return filepath
    
    def _add_executive_summary(self, doc: Document, company_info: Dict, recommendation: Dict):
        """Add executive summary section"""
        
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        rows_data = [
            ('Company Name', company_info.get('company_name', 'N/A')),
            ('Sector', company_info.get('sector', 'N/A')),
            ('Credit Decision', recommendation.get('decision', 'N/A')),
            ('Credit Score', f"{recommendation.get('credit_score', 0)} ({recommendation.get('rating', 'N/A')})"),
            ('Recommended Amount', f"₹ {recommendation.get('recommended_loan_amount', 0):.2f} Cr"),
            ('Recommended Rate', f"{recommendation.get('recommended_interest_rate', 0):.2f}%")
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
    
    def _add_recommendation_section(self, doc: Document, recommendation: Dict):
        """Add detailed recommendation"""
        
        decision = recommendation.get('decision', 'N/A')
        
        # Decision box
        p = doc.add_paragraph()
        run = p.add_run(f"DECISION: {decision}")
        run.bold = True
        run.font.size = Pt(14)
        
        if decision == 'APPROVED':
            run.font.color.rgb = RGBColor(0, 128, 0)
        else:
            run.font.color.rgb = RGBColor(255, 0, 0)
        
        doc.add_paragraph()
        
        # Explanation
        doc.add_paragraph(recommendation.get('explanation', ''))
        
        # Loan details if approved
        if decision == 'APPROVED':
            doc.add_paragraph(f"Recommended Loan Amount: ₹ {recommendation.get('recommended_loan_amount', 0):.2f} Crores")
            doc.add_paragraph(f"Recommended Interest Rate: {recommendation.get('recommended_interest_rate', 0):.2f}% p.a.")
    
    def _add_five_cs_analysis(self, doc: Document, recommendation: Dict):
        """Add Five Cs breakdown"""
        
        five_cs = recommendation.get('five_cs_breakdown', {})
        
        table = doc.add_table(rows=6, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header
        table.rows[0].cells[0].text = 'Criteria'
        table.rows[0].cells[1].text = 'Score'
        table.rows[0].cells[2].text = 'Assessment'
        
        criteria = [
            ('Character', five_cs.get('character', 0)),
            ('Capacity', five_cs.get('capacity', 0)),
            ('Capital', five_cs.get('capital', 0)),
            ('Collateral', five_cs.get('collateral', 0)),
            ('Conditions', five_cs.get('conditions', 0))
        ]
        
        for i, (name, score) in enumerate(criteria, 1):
            table.rows[i].cells[0].text = name
            table.rows[i].cells[1].text = f"{score}/100"
            table.rows[i].cells[2].text = self._get_assessment(score)
    
    def _get_assessment(self, score: float) -> str:
        """Get assessment label for score"""
        if score >= 80:
            return 'Excellent'
        elif score >= 70:
            return 'Good'
        elif score >= 60:
            return 'Satisfactory'
        elif score >= 50:
            return 'Fair'
        else:
            return 'Poor'
    
    def _add_company_overview(self, doc: Document, company_info: Dict, data_summary: Dict):
        """Add company overview section"""
        
        doc.add_paragraph(f"Industry: {company_info.get('sector', 'N/A')}")
        doc.add_paragraph(f"Incorporation Year: {company_info.get('incorporation_year', 'N/A')}")
        doc.add_paragraph(f"CIN: {company_info.get('cin', 'N/A')}")
        doc.add_paragraph(f"GSTIN: {data_summary.get('gst', {}).get('gstin', 'N/A')}")
        
        # Promoters
        promoters = company_info.get('promoters', [])
        if promoters:
            doc.add_paragraph("Key Promoters:")
            for promoter in promoters:
                doc.add_paragraph(f"  • {promoter}", style='List Bullet')
    
    def _add_financial_analysis(self, doc: Document, data_summary: Dict):
        """Add financial analysis section"""
        
        financials = data_summary.get('financials', {})
        
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Light Grid Accent 1'
        
        rows_data = [
            ('Revenue', f"₹ {financials.get('revenue', 0):.2f} Cr"),
            ('EBITDA', f"₹ {financials.get('ebitda', 0):.2f} Cr"),
            ('Net Profit', f"₹ {financials.get('net_profit', 0):.2f} Cr"),
            ('Total Debt', f"₹ {financials.get('total_debt', 0):.2f} Cr"),
            ('Net Worth', f"₹ {financials.get('net_worth', 0):.2f} Cr"),
            ('EBITDA Margin', f"{financials.get('ebitda_margin', 0):.2f}%"),
            ('Debt/EBITDA', f"{financials.get('debt_to_ebitda', 0):.2f}x"),
            ('Debt/Equity', f"{financials.get('debt_to_equity', 0):.2f}x")
        ]
        
        for i, (label, value) in enumerate(rows_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
    
    def _add_gst_banking_analysis(self, doc: Document, data_summary: Dict):
        """Add GST and banking analysis"""
        
        gst = data_summary.get('gst', {})
        bank = data_summary.get('bank', {})
        
        doc.add_heading('GST Analysis', 3)
        doc.add_paragraph(f"Total Turnover (GST): ₹ {gst.get('total_turnover', 0):.2f} Cr")
        doc.add_paragraph(f"Filing Regularity: {gst.get('gst_filing_regularity', 0)} months")
        
        if gst.get('gstr3b_vs_gstr2a_mismatch'):
            p = doc.add_paragraph("⚠ GSTR-3B vs GSTR-2A Mismatch Detected")
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        
        doc.add_heading('Banking Analysis', 3)
        doc.add_paragraph(f"Average Bank Balance: ₹ {bank.get('average_balance', 0):.2f} Lakhs")
        doc.add_paragraph(f"Total Credits: ₹ {bank.get('total_credits', 0):.2f} Lakhs")
        doc.add_paragraph(f"Bounce Rate: {bank.get('bounce_rate', 0)*100:.2f}%")
        
        if bank.get('circular_trading_detected'):
            p = doc.add_paragraph("⚠ Potential Circular Trading Detected")
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    def _add_research_findings(self, doc: Document, research_summary: Dict):
        """Add secondary research findings"""
        
        # News analysis
        news = research_summary.get('news_analysis', [])
        negative_news = [n for n in news if n.get('sentiment') == 'NEGATIVE']
        
        doc.add_paragraph(f"Total News Articles Analyzed: {len(news)}")
        doc.add_paragraph(f"Negative News Count: {len(negative_news)}")
        
        if negative_news:
            doc.add_paragraph("Key Negative News:")
            for article in negative_news[:5]:
                doc.add_paragraph(f"  • {article.get('title', 'N/A')}", style='List Bullet')
        
        # Promoter analysis
        promoter_analysis = research_summary.get('promoter_analysis', {})
        high_risk_promoters = [p for p, data in promoter_analysis.items() 
                              if data.get('risk_level') == 'HIGH']
        
        if high_risk_promoters:
            p = doc.add_paragraph(f"⚠ High Risk Promoters: {', '.join(high_risk_promoters)}")
            p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        
        # Litigation
        litigation = research_summary.get('litigation', {})
        doc.add_paragraph(f"Litigation Cases: {litigation.get('total_cases', 0)}")
    
    def _add_primary_insights(self, doc: Document, primary_insights: Dict):
        """Add primary due diligence insights"""
        
        doc.add_paragraph(f"Capacity Utilization: {primary_insights.get('capacity_utilization', 'N/A')}%")
        doc.add_paragraph(f"Management Quality Score: {primary_insights.get('management_quality_score', 'N/A')}/10")
        doc.add_paragraph(f"Site Visit Score: {primary_insights.get('site_visit_score', 'N/A')}/10")
        
        if primary_insights.get('site_visit_notes'):
            doc.add_paragraph("Site Visit Observations:")
            doc.add_paragraph(primary_insights['site_visit_notes'])
        
        if primary_insights.get('management_interview_notes'):
            doc.add_paragraph("Management Interview Notes:")
            doc.add_paragraph(primary_insights['management_interview_notes'])
    
    def _add_risk_assessment(self, doc: Document, recommendation: Dict):
        """Add risk assessment section"""
        
        risk_factors = recommendation.get('risk_factors', [])
        strengths = recommendation.get('strengths', [])
        
        if risk_factors:
            doc.add_heading('Risk Factors', 3)
            for risk in risk_factors:
                doc.add_paragraph(f"  • {risk}", style='List Bullet')
        
        if strengths:
            doc.add_heading('Strengths', 3)
            for strength in strengths:
                doc.add_paragraph(f"  • {strength}", style='List Bullet')
    
    def _add_terms_conditions(self, doc: Document, recommendation: Dict):
        """Add proposed terms and conditions"""
        
        if recommendation.get('decision') == 'APPROVED':
            doc.add_paragraph(f"Loan Amount: ₹ {recommendation.get('recommended_loan_amount', 0):.2f} Crores")
            doc.add_paragraph(f"Interest Rate: {recommendation.get('recommended_interest_rate', 0):.2f}% p.a.")
            doc.add_paragraph("Tenure: 5 years (indicative)")
            doc.add_paragraph("Repayment: Quarterly installments")
            doc.add_paragraph("Security: Primary and collateral security as per policy")
            doc.add_paragraph("Covenants:")
            doc.add_paragraph("  • Maintain Debt/EBITDA < 4x", style='List Bullet')
            doc.add_paragraph("  • Maintain DSCR > 1.5x", style='List Bullet')
            doc.add_paragraph("  • Quarterly financial reporting", style='List Bullet')
            doc.add_paragraph("  • No dividend distribution without lender consent", style='List Bullet')
        else:
            doc.add_paragraph("Credit facility not recommended at this time.")
            doc.add_paragraph("Reasons outlined in the Risk Assessment section above.")
